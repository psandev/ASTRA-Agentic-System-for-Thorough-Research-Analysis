"""
ASTRA Layer 6 — Report Writer & Visual Generator
Action paradigm: CODE-ACTION (charts, tables, pdf assembly)

Tools:
  write_section, generate_chart_plotly, generate_chart_matplotlib,
  render_mermaid, build_pdf, build_markdown, embed_figure
"""
from __future__ import annotations

import base64
import json
import os
import re
import subprocess
from pathlib import Path
from typing import Any, Optional

from langchain_core.tools import tool
from loguru import logger
from openai import OpenAI

from astra.config import get_config

# ─── write_section ───────────────────────────────────────────────────────────

_WRITER_SYSTEM = """You are ASTRA's expert research writer.

Write a comprehensive, well-structured report section in academic markdown, formatted like a high-quality scientific paper or detailed Medium article.

Requirements:
- Use headers (## for main, ### for sub-sections)
- Cite sources inline as [Source N] immediately after the claim they support
- Include all relevant data, statistics, and findings from the retrieved chunks
- Generate Markdown tables for comparative/benchmark data (use | col | col | format)
  * Tables MUST include at least one column with numeric values (scores, latency ms, %, ranks, counts)
    so that a companion chart can be auto-generated. If raw numbers are unavailable, use ordinal
    scores (e.g. 1–5 scale) or derive approximate percentages.
  * Place every table IMMEDIATELY after the paragraph that introduces it — never cluster tables
    at the end of the section.
- For mathematical formulas use LaTeX syntax: inline as $formula$ and block as $$formula$$
  Example: The BM25 score is $score(D,Q) = \\sum_{i=1}^{n} IDF(q_i) \\cdot \\frac{f(q_i, D) \\cdot (k_1 + 1)}{f(q_i, D) + k_1 \\cdot (1 - b + b \\cdot \\frac{|D|}{avgdl})}$
- For hyperlinks use Markdown syntax: [link text](URL)
- Be analytical, not just descriptive
- Minimum 300 words per section
- Do NOT add a "References", "Key Sources", or "Further Reading" subsection — all citations are
  consolidated in the global Bibliography at the end of the report.

Return ONLY the markdown content for the section."""


@tool
def write_section(
    section_title: str,
    section_outline: str,
    retrieved_chunks: list[dict],
    style: str = "academic",
) -> dict:
    """
    Layer 6: Generate a report section with inline citations and analytical prose.
    Uses Qwen3-32B-AWQ at temperature=0.7.

    Args:
        section_title: Title of the section.
        section_outline: Brief outline of what this section should cover.
        retrieved_chunks: List of {text, metadata} dicts from RAG.
        style: Writing style: "academic" | "technical" | "executive".

    Returns:
        {markdown, citations, word_count}
    """
    logger.info(f"[Layer 6] Writing section: '{section_title}'")
    cfg = get_config()

    # Format retrieved chunks as numbered sources
    sources_text = ""
    citations: list[dict] = []
    for i, chunk in enumerate(retrieved_chunks[:15], start=1):
        text = chunk.get("text", "")
        meta = chunk.get("metadata", {})
        url = meta.get("source_url", "")
        sources_text += f"\n[Source {i}] {url}\n{text[:600]}\n"
        citations.append(
            {
                "key": f"Source {i}",
                "url": url,
                "text_preview": text[:100],
            }
        )

    user_msg = f"""Write the report section '{section_title}'.

Section outline: {section_outline}

Style: {style}

Retrieved sources (cite as [Source N]):
{sources_text[:6000]}

Write the full section now:"""

    try:
        client = OpenAI(base_url=cfg.vllm_base_url, api_key=cfg.vllm_api_key)
        response = client.chat.completions.create(
            model=cfg.astra_main_model,
            messages=[
                {"role": "system", "content": _WRITER_SYSTEM},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.7,
            max_tokens=cfg.astra_main_max_tokens,
        )
        markdown = response.choices[0].message.content or ""
        word_count = len(markdown.split())

        logger.info(f"[Layer 6] '{section_title}' written: {word_count} words")
        return {
            "markdown": markdown,
            "citations": citations,
            "word_count": word_count,
        }

    except Exception as e:
        logger.error(f"[Layer 6] write_section failed: {e}")
        # Return a stub section
        stub = (
            f"## {section_title}\n\n"
            f"*Section generation failed: {str(e)}. "
            f"Manual review required.*\n\n"
            + "\n\n".join(c["text_preview"] for c in citations[:5])
        )
        return {
            "markdown": stub,
            "citations": citations,
            "word_count": len(stub.split()),
        }


# ─── generate_chart_plotly ────────────────────────────────────────────────────

@tool
def generate_chart_plotly(
    chart_type: str,
    data: dict,
    title: str,
    x_label: str = "",
    y_label: str = "",
    output_path: str = "",
    theme: str = "plotly_white",
) -> dict:
    """
    Layer 6: Create interactive charts using Plotly. Preferred for HTML output.

    Args:
        chart_type: "bar" | "line" | "scatter" | "pie" | "heatmap" | "box".
        data: {x: [...], y: [...]} or {labels: [...], values: [...]}.
        title: Chart title.
        x_label: X-axis label.
        y_label: Y-axis label.
        output_path: Base path for output (without extension).
        theme: Plotly theme name.

    Returns:
        {html_path, png_path}
    """
    import plotly.graph_objects as go
    import plotly.express as px

    cfg = get_config()
    if not output_path:
        safe_title = re.sub(r"[^\w\-_]", "_", title)[:50]
        output_path = str(cfg.get_chart_path(safe_title))

    logger.info(f"[Layer 6] Plotly chart: {chart_type} '{title}'")

    try:
        x_data = data.get("x", data.get("labels", []))
        y_data = data.get("y", data.get("values", []))

        if chart_type == "bar":
            fig = px.bar(x=x_data, y=y_data, title=title, template=theme)
        elif chart_type == "line":
            fig = px.line(x=x_data, y=y_data, title=title, template=theme)
        elif chart_type == "scatter":
            fig = px.scatter(x=x_data, y=y_data, title=title, template=theme)
        elif chart_type == "pie":
            fig = px.pie(names=x_data, values=y_data, title=title, template=theme)
        elif chart_type == "heatmap":
            z_data = data.get("z", [])
            fig = px.imshow(z_data, title=title, template=theme)
        else:
            fig = px.bar(x=x_data, y=y_data, title=title, template=theme)

        # Add axis labels
        fig.update_layout(
            xaxis_title=x_label,
            yaxis_title=y_label,
            title={"text": title, "x": 0.5},
        )

        # Save outputs
        html_path = f"{output_path}.html"
        png_path = f"{output_path}.png"

        fig.write_html(html_path)
        try:
            fig.write_image(png_path, width=1200, height=600)
        except Exception:
            logger.warning("[Layer 6] Kaleido not available — skipping PNG export")
            png_path = ""

        logger.info(f"[Layer 6] Chart saved: {html_path}")
        return {"html_path": html_path, "png_path": png_path}

    except Exception as e:
        logger.error(f"[Layer 6] Plotly chart failed: {e}")
        return {"html_path": "", "png_path": ""}


# ─── generate_chart_matplotlib ────────────────────────────────────────────────

@tool
def generate_chart_matplotlib(
    chart_type: str,
    data: dict,
    title: str,
    output_path: str = "",
    dpi: int = 150,
    style: str = "seaborn-v0_8-darkgrid",
) -> dict:
    """
    Layer 6: Create static publication-quality charts using Matplotlib.

    Args:
        chart_type: "bar" | "line" | "scatter" | "hist" | "heatmap".
        data: {x: [...], y: [...], labels: [...]}.
        title: Chart title.
        output_path: Base output path (without extension).
        dpi: Image resolution.
        style: Matplotlib style.

    Returns:
        {png_path, svg_path}
    """
    import matplotlib
    matplotlib.use("Agg")  # Non-interactive backend
    import matplotlib.pyplot as plt

    cfg = get_config()
    if not output_path:
        safe_title = re.sub(r"[^\w\-_]", "_", title)[:50]
        output_path = str(cfg.get_chart_path(safe_title))

    logger.info(f"[Layer 6] Matplotlib chart: {chart_type} '{title}'")

    try:
        try:
            plt.style.use(style)
        except Exception:
            plt.style.use("default")

        fig, ax = plt.subplots(figsize=(10, 6))

        x_data = data.get("x", data.get("labels", []))
        y_data = data.get("y", data.get("values", []))

        if chart_type == "bar":
            ax.bar(range(len(x_data)), y_data, tick_label=x_data)
            ax.set_xlabel(data.get("x_label", ""))
            ax.set_ylabel(data.get("y_label", ""))
            plt.xticks(rotation=45, ha="right")
        elif chart_type == "line":
            ax.plot(x_data, y_data, marker="o")
        elif chart_type == "scatter":
            ax.scatter(x_data, y_data)
        elif chart_type == "hist":
            ax.hist(y_data, bins=data.get("bins", 20))
        elif chart_type == "heatmap":
            import numpy as np
            z = np.array(data.get("z", [[0]]))
            im = ax.imshow(z, aspect="auto", cmap="viridis")
            fig.colorbar(im, ax=ax)
        else:
            ax.bar(range(len(x_data)), y_data, tick_label=x_data)

        ax.set_title(title, fontsize=14, fontweight="bold")
        ax.set_xlabel(data.get("x_label", ""), fontsize=11)
        ax.set_ylabel(data.get("y_label", ""), fontsize=11)
        plt.tight_layout()

        png_path = f"{output_path}.png"
        svg_path = f"{output_path}.svg"
        fig.savefig(png_path, dpi=dpi, bbox_inches="tight")
        fig.savefig(svg_path, format="svg", bbox_inches="tight")
        plt.close(fig)

        logger.info(f"[Layer 6] Chart saved: {png_path}")
        return {"png_path": png_path, "svg_path": svg_path}

    except Exception as e:
        logger.error(f"[Layer 6] Matplotlib chart failed: {e}")
        return {"png_path": "", "svg_path": ""}


# ─── render_mermaid ───────────────────────────────────────────────────────────

@tool
def render_mermaid(
    mermaid_code: str,
    output_path: str = "",
    theme: str = "default",
) -> dict:
    """
    Layer 6: Render Mermaid.js diagrams (flowcharts, sequence, ER, class).
    Requires `npm install -g @mermaid-js/mermaid-cli` (mmdc).

    Args:
        mermaid_code: Mermaid diagram source code.
        output_path: Output file path (without extension).
        theme: "default" | "dark" | "forest" | "neutral".

    Returns:
        {png_path, svg_path}
    """
    cfg = get_config()
    if not output_path:
        import hashlib
        h = hashlib.md5(mermaid_code.encode()).hexdigest()[:8]
        output_path = str(cfg.get_chart_path(f"mermaid_{h}"))

    svg_path = f"{output_path}.svg"
    png_path = f"{output_path}.png"

    try:
        # Write mermaid source to temp file
        mmd_path = f"{output_path}.mmd"
        Path(mmd_path).write_text(mermaid_code, encoding="utf-8")

        # Try mmdc (mermaid CLI)
        result = subprocess.run(
            ["mmdc", "-i", mmd_path, "-o", svg_path, "-t", theme, "--quiet"],
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode != 0:
            raise RuntimeError(result.stderr)

        # Also export PNG
        subprocess.run(
            ["mmdc", "-i", mmd_path, "-o", png_path, "-t", theme, "-w", "1200", "--quiet"],
            capture_output=True,
            timeout=30,
        )

        logger.info(f"[Layer 6] Mermaid diagram rendered: {svg_path}")
        return {"png_path": png_path, "svg_path": svg_path}

    except FileNotFoundError:
        logger.warning("[Layer 6] mmdc not found — saving mermaid source as .mmd only")
        return {"png_path": "", "svg_path": f"{output_path}.mmd"}
    except Exception as e:
        logger.error(f"[Layer 6] Mermaid render failed: {e}")
        return {"png_path": "", "svg_path": ""}


# ─── build_docx ───────────────────────────────────────────────────────────────

@tool
def build_docx(
    sections: list[dict],
    figures: list[dict],
    bibliography: list[dict],
    template_path: str = "",
    output_path: str = "",
    include_toc: bool = True,
) -> dict:
    """
    Layer 6: Assemble the final .docx report from sections, figures, tables.
    Uses python-docx with ASTRA template.

    Args:
        sections: List of {title, markdown, citations} dicts.
        figures: List of {path, caption} dicts.
        bibliography: List of {key, title, authors, year, url} dicts.
        template_path: Path to .docx template (optional).
        output_path: Output file path.
        include_toc: Include table of contents placeholder.

    Returns:
        {docx_path, page_count, word_count}
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cfg = get_config()
    if not output_path:
        import time
        ts = int(time.time())
        output_path = str(cfg.get_output_path(f"ASTRA_Report_{ts}.docx"))

    logger.info(f"[Layer 6] Building DOCX: {output_path}")

    try:
        # Use template if available, else create new document
        template = Path(template_path or cfg.astra_docx_template)
        if template.exists():
            doc = Document(str(template))
        else:
            doc = Document()
            # Set default style
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

        # Title page
        title_para = doc.add_heading("ASTRA Research Report", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if include_toc:
            doc.add_heading("Table of Contents", level=1)
            doc.add_paragraph("[TOC — update field in Word]")
            doc.add_page_break()

        total_words = 0
        for sec in sections:
            title = sec.get("title", "")
            content = sec.get("markdown", "")

            # Add section heading
            doc.add_heading(title, level=1)

            # Process markdown → docx (basic conversion)
            _markdown_to_docx(doc, content)
            total_words += len(content.split())

            # Embed charts/figures for this section
            chart_paths = sec.get("chart_paths", [])
            for cp in chart_paths:
                if cp and Path(cp).exists() and cp.endswith(".png"):
                    doc.add_picture(cp, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Figures appendix
        if figures:
            doc.add_heading("Figures", level=1)
            for fig in figures:
                fp = fig.get("path", "")
                if fp and Path(fp).exists():
                    doc.add_picture(fp, width=Inches(5.0))
                    cap_para = doc.add_paragraph(fig.get("caption", ""))
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Bibliography
        doc.add_page_break()
        doc.add_heading("Bibliography", level=1)
        for bib in bibliography:
            authors = ", ".join(bib.get("authors", [bib.get("key", "")]))
            year = bib.get("year", "")
            title_str = bib.get("title", "")
            url = bib.get("url", "")
            entry = f"{authors} ({year}). {title_str}. {url}"
            doc.add_paragraph(entry, style="List Bullet")

        doc.save(output_path)

        # Estimate page count (rough: 250 words/page)
        page_count = max(1, total_words // 250)

        logger.info(
            f"[Layer 6] DOCX built: {output_path} (~{page_count}p, {total_words}w)"
        )
        return {
            "docx_path": output_path,
            "page_count": page_count,
            "word_count": total_words,
        }

    except Exception as e:
        logger.error(f"[Layer 6] build_docx failed: {e}")
        return {"docx_path": "", "page_count": 0, "word_count": 0, "error": str(e)}


def _markdown_to_docx(doc, markdown: str) -> None:
    """Basic markdown → python-docx converter (handles headers, bold, paragraphs)."""
    from docx.shared import Pt

    for line in markdown.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("| ") or stripped.startswith("|---"):
            # Skip table lines (tables need proper rendering)
            pass
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        else:
            para = doc.add_paragraph()
            # Handle basic bold **text**
            parts = re.split(r"\*\*(.+?)\*\*", stripped)
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True


# ─── build_pdf ────────────────────────────────────────────────────────────────

_PDF_CSS = """
@page {
    size: A4;
    margin: 2.5cm 2cm 2.5cm 2cm;
    @bottom-center { content: "Page " counter(page) " of " counter(pages); font-size: 9pt; color: #666; }
}
body        { font-family: "DejaVu Serif", Georgia, serif; font-size: 11pt; line-height: 1.6; color: #1a1a1a; }
h1          { font-size: 22pt; color: #003366; border-bottom: 2px solid #003366; padding-bottom: 6px; margin-top: 0; }
h2          { font-size: 16pt; color: #003366; margin-top: 28px; }
h3          { font-size: 13pt; color: #1a5276; }
p           { margin: 0.5em 0 0.8em 0; text-align: justify; }
ul, ol      { margin: 0.4em 0 0.8em 1.4em; }
li          { margin-bottom: 0.3em; }
table       { border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 10pt; }
th          { background: #003366; color: #fff; padding: 6px 10px; text-align: left; }
td          { border: 1px solid #ccc; padding: 5px 10px; }
tr:nth-child(even) td { background: #f5f8fc; }
a           { color: #0066cc; text-decoration: underline; }
code        { font-family: "DejaVu Sans Mono", monospace; font-size: 9pt; background: #f4f4f4; padding: 1px 4px; border-radius: 3px; }
pre         { background: #f4f4f4; padding: 10px; border-left: 4px solid #003366; overflow-x: auto; font-size: 9pt; }
blockquote  { border-left: 4px solid #aaa; margin: 0.5em 0; padding: 0.4em 1em; color: #555; background: #fafafa; }
img         { max-width: 100%; height: auto; display: block; margin: 1em auto; }
.title-page { text-align: center; padding-top: 80px; page-break-after: always; }
.title-page h1 { border: none; font-size: 28pt; margin-bottom: 12px; }
.title-page .subtitle { font-size: 14pt; color: #555; }
.toc        { page-break-after: always; }
.toc a      { color: #003366; text-decoration: none; }
.section    { page-break-inside: avoid; }
hr          { border: none; border-top: 1px solid #ccc; margin: 2em 0; }
"""


def _markdown_to_html(md: str) -> str:
    """
    Convert markdown to HTML with:
    - Active hyperlinks ([text](url) → <a href>)
    - LaTeX math preserved for MathJax ($...$ and $$...$$)
    - Tables, fenced code blocks, bold/italic
    Uses markdown2 when available, falls back to inline regex converter.
    """
    # Step 1: Protect block math ($$...$$) from markdown processing
    # Replace with unique placeholders so markdown2 won't mangle them
    _math_store: dict[str, str] = {}
    _counter = [0]

    def _save(content: str, display: bool) -> str:
        key = f"XMATHX{_counter[0]}X"
        _counter[0] += 1
        if display:
            _math_store[key] = f"\\[{content}\\]"
        else:
            _math_store[key] = f"\\({content}\\)"
        return key

    # Block math: $$...$$
    md2 = re.sub(
        r"\$\$(.+?)\$\$",
        lambda m: _save(m.group(1).strip(), display=True),
        md,
        flags=re.DOTALL,
    )
    # Inline math: $...$ (not followed by another $)
    md2 = re.sub(
        r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)",
        lambda m: _save(m.group(1).strip(), display=False),
        md2,
    )
    # Also protect \[...\] display math from markdown2 (LLM often writes this style)
    md2 = re.sub(
        r"\\\[(.+?)\\\]",
        lambda m: _save(m.group(1).strip(), display=True),
        md2,
        flags=re.DOTALL,
    )
    # Also protect \(...\) inline math from markdown2 (LLM often writes this style)
    md2 = re.sub(
        r"\\\((.+?)\\\)",
        lambda m: _save(m.group(1).strip(), display=False),
        md2,
    )

    try:
        import markdown2
        html = markdown2.markdown(
            md2,
            extras=[
                "tables",
                "fenced-code-blocks",
                "strike",
                "footnotes",
                "header-ids",
                "break-on-newline",
            ],
        )
    except ImportError:
        html = _basic_md_to_html(md2)

    # Step 2: Restore math placeholders
    for key, value in _math_store.items():
        html = html.replace(key, value)

    return html


def _latex_to_png_base64(formula: str, display: bool) -> str:
    """
    Render a LaTeX formula to a base64-encoded PNG using matplotlib's mathtext
    engine (no external LaTeX installation required).

    Used by _render_math_for_pdf() to replace MathJax delimiters with static
    images before handing HTML to weasyprint (which cannot execute JavaScript).
    Returns "" on failure so callers can fall back to raw text.
    """
    try:
        import io
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fontsize = 15 if display else 11
        fig = plt.figure(figsize=(0.01, 0.01))
        text = fig.text(0, 0, f"${formula}$", fontsize=fontsize)
        fig.canvas.draw()
        renderer = fig.canvas.get_renderer()
        bbox = text.get_window_extent(renderer=renderer)
        w_in = max(0.5, bbox.width / fig.dpi + 0.15)
        h_in = max(0.25, bbox.height / fig.dpi + 0.10)
        fig.set_size_inches(w_in, h_in)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                    transparent=True, pad_inches=0.04)
        plt.close(fig)
        buf.seek(0)
        return base64.b64encode(buf.read()).decode()
    except Exception:
        return ""


def _render_math_for_pdf(html: str) -> str:
    """
    Replace MathJax-style LaTeX delimiters with rendered PNG <img> tags.

    Converts:
      \\[...\\]  →  <div><img src="data:image/png;base64,..."></div>
      \\(...\\)  →  <img src="data:image/png;base64,..." style="vertical-align:middle">

    Falls back to <code>formula</code> when matplotlib mathtext cannot parse
    the expression (e.g. unsupported commands). Called inside build_pdf() before
    the HTML is passed to weasyprint so formulas render as images rather than
    as raw LaTeX markup.
    """
    def _sub_display(m: re.Match) -> str:
        formula = m.group(1).strip()
        b64 = _latex_to_png_base64(formula, display=True)
        if b64:
            return (
                f'<div style="text-align:center;margin:0.8em 0">'
                f'<img src="data:image/png;base64,{b64}" style="max-width:95%">'
                f"</div>"
            )
        return (
            f'<div style="text-align:center;font-family:monospace;'
            f'margin:0.8em 0;font-size:9pt">{formula}</div>'
        )

    def _sub_inline(m: re.Match) -> str:
        formula = m.group(1).strip()
        b64 = _latex_to_png_base64(formula, display=False)
        if b64:
            return (
                f'<img src="data:image/png;base64,{b64}" '
                f'style="vertical-align:middle;max-height:1.4em">'
            )
        return f"<code>{formula}</code>"

    html = re.sub(r"\\\[(.+?)\\\]", _sub_display, html, flags=re.DOTALL)
    html = re.sub(r"\\\((.+?)\\\)", _sub_inline, html)
    return html


def _best_injection_point(section_html: str, caption: str) -> int:
    """Return the character position (end of a </p> or </table> tag) inside
    *section_html* where a figure with the given *caption* should be inserted.

    Strategy:
      1. Extract content keywords from caption (words ≥ 4 chars).
      2. Score every <p>…</p> paragraph by keyword overlap.
      3. Inject after the highest-scoring paragraph.
      4. If no paragraph scores > 0, inject after the LAST </p> in the section
         (before any trailing headings/lists) so it stays within the narrative.
    """
    keywords = {w.lower() for w in re.findall(r"[a-zA-Z]{4,}", caption)}

    best_score = 0
    best_end = -1

    for m in re.finditer(r"<p[^>]*>(.*?)</p>", section_html, re.DOTALL):
        para_plain = re.sub(r"<[^>]+>", "", m.group(1)).lower()
        para_words = {w for w in re.findall(r"[a-zA-Z]{4,}", para_plain)}
        score = len(keywords & para_words)
        if score > best_score:
            best_score = score
            best_end = m.end()

    if best_end >= 0:
        return best_end

    # Fallback: after the last </p>
    last_p = section_html.rfind("</p>")
    if last_p >= 0:
        return last_p + len("</p>")

    return len(section_html)


def _strip_section_references(html: str) -> str:
    """Remove per-section 'References' / 'Key Sources' subsections that the
    LLM sometimes appends despite instructions.  These are redundant with the
    global Bibliography and disrupt narrative flow.

    Matches any <h3> or <h4> whose text is References / Key Sources / Further
    Reading (case-insensitive) and removes that heading + the immediately
    following block (<ul>, <ol>, or <p> with source links).
    """
    # Match heading + following list or paragraph block
    html = re.sub(
        r"<h[34][^>]*>\s*(?:References|Key\s+Sources|Further\s+Reading|Sources)\s*</h[34]>"
        r"\s*(?:<[ou]l>.*?</[ou]l>|<p>.*?</p>)",
        "",
        html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return html


def _linkify_bare_urls(html: str) -> str:
    """Convert bare (non-linked) URLs in HTML text nodes to active hyperlinks.

    Splits on HTML tags and tracks anchor-tag depth so URLs that already appear
    as both the href attribute value AND link text are never double-wrapped.
    """
    parts = re.split(r"(<[^>]+>)", html)
    result: list[str] = []
    in_anchor = 0  # track <a> nesting depth

    for part in parts:
        if part.startswith("<"):
            low = part.lower()
            # Detect opening / closing anchor tags
            if re.match(r"<a[\s>]", low):
                in_anchor += 1
            elif low.startswith("</a"):
                in_anchor = max(0, in_anchor - 1)
            result.append(part)
        else:
            if in_anchor == 0:
                # Safe text node — linkify bare URLs
                part = re.sub(
                    r"(https?://[^\s<>\"',\)\]]+)",
                    r'<a href="\1" target="_blank">\1</a>',
                    part,
                )
            result.append(part)

    return "".join(result)


def _basic_md_to_html(md: str) -> str:
    """Regex-based fallback markdown→HTML converter with link + table support."""
    lines = md.split("\n")
    out: list[str] = []
    in_ul = False
    in_ol = False
    in_code = False
    in_table = False
    table_header_done = False

    def _inline(text: str) -> str:
        """Apply inline formatting: bold, italic, code, links."""
        # Links [text](url)
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', text)
        # Bold **text**
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        # Italic *text*
        text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
        # Inline code `code`
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        return text

    def _close_lists() -> None:
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    def _close_table() -> None:
        nonlocal in_table, table_header_done
        if in_table:
            out.append("</tbody></table>")
            in_table = False
            table_header_done = False

    for line in lines:
        # Fenced code blocks
        if line.strip().startswith("```"):
            if in_code:
                out.append("</code></pre>")
                in_code = False
            else:
                _close_lists()
                _close_table()
                lang = line.strip()[3:].strip()
                cls = f' class="language-{lang}"' if lang else ""
                out.append(f"<pre><code{cls}>")
                in_code = True
            continue
        if in_code:
            out.append(line.replace("&", "&amp;").replace("<", "&lt;"))
            continue

        stripped = line.strip()

        # Markdown table rows
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c.replace(" ", "")) for c in cells):
                # Separator row — switch to tbody
                _close_lists()
                out.append("</thead><tbody>")
                table_header_done = True
                continue
            if not in_table:
                _close_lists()
                out.append('<table><thead>')
                in_table = True
                table_header_done = False
                tag = "th"
            else:
                tag = "td"
            row_html = "".join(f"<{tag}>{_inline(c)}</{tag}>" for c in cells)
            out.append(f"<tr>{row_html}</tr>")
            continue
        else:
            _close_table()

        if not stripped:
            _close_lists()
            continue

        if stripped.startswith("### "):
            _close_lists()
            out.append(f"<h3>{_inline(stripped[4:])}</h3>")
        elif stripped.startswith("## "):
            _close_lists()
            out.append(f"<h2>{_inline(stripped[3:])}</h2>")
        elif stripped.startswith("# "):
            _close_lists()
            out.append(f"<h2>{_inline(stripped[2:])}</h2>")
        elif stripped.startswith(("- ", "* ")):
            if not in_ul:
                _close_lists()
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{_inline(stripped[2:])}</li>")
        elif re.match(r"^\d+\. ", stripped):
            if not in_ol:
                _close_lists()
                out.append("<ol>")
                in_ol = True
            out.append(f"<li>{_inline(re.sub(r'^\d+\. ', '', stripped))}</li>")
        elif re.match(r"^-{3,}$|^\*{3,}$", stripped):
            _close_lists()
            out.append("<hr>")
        elif stripped.startswith("> "):
            _close_lists()
            out.append(f"<blockquote><p>{_inline(stripped[2:])}</p></blockquote>")
        else:
            _close_lists()
            out.append(f"<p>{_inline(stripped)}</p>")

    _close_lists()
    _close_table()
    if in_code:
        out.append("</code></pre>")

    return "\n".join(out)


def _embed_image(img_path: str) -> str:
    """Embed image as a base64 data URI (used by both build_html and build_pdf)."""
    try:
        p = Path(img_path)
        if not p.exists():
            return img_path
        ext = p.suffix.lower().lstrip(".")
        mime = {
            "png": "image/png",
            "jpg": "image/jpeg",
            "jpeg": "image/jpeg",
            "svg": "image/svg+xml",
        }.get(ext, "image/png")
        b64 = base64.b64encode(p.read_bytes()).decode()
        return f"data:{mime};base64,{b64}"
    except Exception:
        return img_path


@tool
def build_pdf(
    sections: list[dict],
    figures: list[dict],
    bibliography: list[dict],
    output_path: str = "",
    include_toc: bool = True,
) -> dict:
    """
    Layer 6: Assemble the final PDF report from sections, figures, and bibliography.
    Uses weasyprint for professional HTML→PDF conversion.

    Args:
        sections: List of {title, markdown, citations, chart_paths} dicts.
        figures: List of {path, caption} dicts.
        bibliography: List of {key, title, authors, year, url} dicts.
        output_path: Output .pdf file path.
        include_toc: Include table of contents.

    Returns:
        {pdf_path, page_count, word_count}
    """
    import time
    from weasyprint import HTML, CSS

    cfg = get_config()
    if not output_path:
        ts = int(time.time())
        output_path = str(cfg.get_output_path(f"ASTRA_Report_{ts}.pdf"))

    logger.info(f"[Layer 6] Building PDF: {output_path}")

    try:
        parts: list[str] = []

        # Title page
        parts.append(
            '<div class="title-page">'
            "<h1>ASTRA Research Report</h1>"
            '<p class="subtitle">Generated by ASTRA · Multi-Agent Research Pipeline</p>'
            "</div>"
        )

        # Table of contents
        if include_toc:
            toc_items = "".join(
                f'<li><a href="#sec-{i}">{sec.get("title", f"Section {i}")}</a></li>'
                for i, sec in enumerate(sections, 1)
            )
            parts.append(
                '<div class="toc">'
                "<h1>Table of Contents</h1>"
                f"<ol>{toc_items}</ol>"
                "</div>"
            )

        total_words = 0

        for i, sec in enumerate(sections, 1):
            title = sec.get("title", f"Section {i}")
            content = sec.get("markdown", "")
            chart_paths = sec.get("chart_paths", [])

            section_html = _render_math_for_pdf(_markdown_to_html(content))
            # Link [Source N] citations to bibliography anchors
            section_html = re.sub(
                r"\[Source (\d+)\]",
                r'<a href="#bib-\1">[Source \1]</a>',
                section_html,
            )
            total_words += len(content.split())

            # Embed chart images as base64 data URIs so weasyprint can render them
            charts_html = ""
            for cp in chart_paths:
                if cp and Path(cp).exists() and cp.endswith(".png"):
                    charts_html += (
                        f'<figure><img src="{_embed_image(cp)}" alt="Chart">'
                        f"<figcaption>Chart</figcaption></figure>"
                    )

            parts.append(
                f'<div class="section" id="sec-{i}">'
                f"<h1>{title}</h1>"
                f"{section_html}"
                f"{charts_html}"
                "</div>"
            )

        # Appendix: additional figures
        if figures:
            parts.append("<h1>Figures</h1>")
            for fig in figures:
                fp = fig.get("path", "")
                if fp and Path(fp).exists():
                    cap = fig.get("caption", "")
                    parts.append(
                        f'<figure><img src="{_embed_image(fp)}" alt="{cap}">'
                        f"<figcaption>{cap}</figcaption></figure>"
                    )

        # Bibliography — id="bib-N" enables [Source N] anchor links from sections
        bib_items = ""
        for j, bib in enumerate(bibliography, 1):
            authors = ", ".join(bib.get("authors", [bib.get("key", "Unknown")]))
            year = bib.get("year", "n.d.")
            title_str = bib.get("title", "")
            url = bib.get("url", "")
            url_link = f'<a href="{url}">{url}</a>' if url else ""
            bib_items += (
                f'<li id="bib-{j}">{authors} ({year}). <em>{title_str}</em>. {url_link}</li>'
            )

        parts.append(f"<h1>Bibliography</h1><ol>{bib_items}</ol>")

        full_html = (
            "<!DOCTYPE html><html><head>"
            '<meta charset="utf-8">'
            f"<style>{_PDF_CSS}</style>"
            "</head><body>"
            + "\n".join(parts)
            + "</body></html>"
        )

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        HTML(string=full_html).write_pdf(
            output_path,
            stylesheets=[CSS(string=_PDF_CSS)],
        )

        page_count = max(1, total_words // 250)
        logger.info(
            f"[Layer 6] PDF built: {output_path} (~{page_count}p, {total_words}w)"
        )
        return {
            "pdf_path": output_path,
            "page_count": page_count,
            "word_count": total_words,
        }

    except Exception as e:
        logger.error(f"[Layer 6] build_pdf failed: {e}")
        return {"pdf_path": "", "page_count": 0, "word_count": 0, "error": str(e)}


# ─── build_markdown ───────────────────────────────────────────────────────────

@tool
def build_markdown(
    sections: list[dict],
    bibliography: list[dict],
    output_path: str = "",
) -> dict:
    """
    Layer 6: Build the final .md report with embedded citations.

    Args:
        sections: List of {title, markdown, citations} dicts.
        bibliography: List of {key, title, authors, year, url} dicts.
        output_path: Output file path.

    Returns:
        {md_path}
    """
    cfg = get_config()
    if not output_path:
        import time
        ts = int(time.time())
        output_path = str(cfg.get_output_path(f"ASTRA_Report_{ts}.md"))

    logger.info(f"[Layer 6] Building Markdown report: {output_path}")

    parts: list[str] = [
        "# ASTRA Research Report",
        "",
        "---",
        "",
        "## Table of Contents",
        "",
    ]

    # TOC
    for i, sec in enumerate(sections, start=1):
        title = sec.get("title", f"Section {i}")
        anchor = title.lower().replace(" ", "-").replace("/", "")
        parts.append(f"{i}. [{title}](#{anchor})")

    parts.extend(["", "---", ""])

    total_words = 0

    for sec in sections:
        title = sec.get("title", "")
        content = sec.get("markdown", "")
        chart_paths = sec.get("chart_paths", [])

        parts.append(f"## {title}")
        parts.append("")
        parts.append(content)
        total_words += len(content.split())

        # Embed chart images
        for cp in chart_paths:
            if cp and Path(cp).exists():
                parts.append(f"\n![Chart]({cp})\n")

        parts.append("")
        parts.append("---")
        parts.append("")

    # Bibliography
    parts.extend(["## Bibliography", ""])
    for bib in bibliography:
        authors = ", ".join(bib.get("authors", [bib.get("key", "Unknown")]))
        year = bib.get("year", "n.d.")
        title_str = bib.get("title", "")
        url = bib.get("url", "")
        parts.append(f"- {authors} ({year}). *{title_str}*. {url}")

    full_md = "\n".join(parts)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(full_md, encoding="utf-8")

    logger.info(f"[Layer 6] Markdown report built: {output_path} ({total_words}w)")
    return {"md_path": output_path, "word_count": total_words}


# ─── build_html ───────────────────────────────────────────────────────────────

_MATHJAX_SCRIPT = """\
<script>
MathJax = {
  tex: {
    inlineMath: [['\\\\(', '\\\\)'], ['$', '$']],
    displayMath: [['\\\\[', '\\\\]'], ['$$', '$$']],
    processEscapes: true
  },
  options: { skipHtmlTags: ['script','noscript','style','textarea','pre'] }
};
</script>
<script defer src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>"""

_HTML_EXTRA_CSS = """
/* ── Base typography ─────────────────────────────────────────────────── */
body        { font-family: "Georgia", "Times New Roman", serif;
              font-size: 17px; line-height: 1.75; color: #1a1a2e; max-width: 900px;
              margin: 0 auto; padding: 2rem 2.5rem; }
/* Section headers */
h1          { font-size: 1.7em; color: #003366; border-bottom: 3px solid #003366;
              padding-bottom: 10px; margin-top: 2.5rem; }
h2          { font-size: 1.35em; color: #004080; margin-top: 2rem; border-bottom: 1px solid #cde; padding-bottom: 4px; }
h3          { font-size: 1.15em; color: #1a5276; margin-top: 1.6rem; }
h4          { font-size: 1em; color: #2e4057; font-style: italic; margin-top: 1.2rem; }
/* Links */
a           { color: #0055cc; text-decoration: none; }
a:hover     { text-decoration: underline; }
a.cite      { font-size: 0.82em; vertical-align: super; color: #0055aa;
              font-family: sans-serif; }
/* Paragraphs */
p           { margin: 0.4em 0 1em; text-align: justify; }
/* ── Tables ──────────────────────────────────────────────────────────── */
table       { border-collapse: collapse; width: 100%; margin: 1.4em 0 0.4em;
              font-size: 0.93em; font-family: sans-serif; }
caption     { font-size: 0.92em; color: #444; font-style: italic;
              caption-side: top; text-align: left; margin-bottom: 4px; }
th          { background: #003366; color: #fff; padding: 9px 13px; text-align: left;
              font-weight: 600; letter-spacing: 0.02em; }
td          { border: 1px solid #c5cfe8; padding: 7px 13px; vertical-align: top; }
tr:nth-child(even) td { background: #f0f4fb; }
tr:hover td { background: #e8eef8; }
/* ── Code ────────────────────────────────────────────────────────────── */
code        { font-family: "SFMono-Regular", Consolas, "Liberation Mono", monospace;
              font-size: 0.85em; background: #f0f4f8; padding: 2px 6px;
              border-radius: 3px; color: #c0392b; }
pre         { background: #1e1e2e; color: #cdd6f4; padding: 16px 20px;
              border-radius: 8px; overflow-x: auto; font-size: 0.88em; }
pre code    { background: none; padding: 0; color: inherit; }
/* ── Block quotes / callouts ─────────────────────────────────────────── */
blockquote  { border-left: 4px solid #0066cc; margin: 1.2em 0;
              padding: 0.6em 1.4em; background: #f5f9ff; color: #333;
              border-radius: 0 6px 6px 0; font-style: italic; }
/* ── Figures & charts — inline, scientific-paper style ──────────────── */
.chart-figure {
    margin: 1.8em auto;
    text-align: center;
    page-break-inside: avoid;
    max-width: 820px;
}
.chart-figure figure {
    display: inline-block;
    width: 100%;
    border: 1px solid #d0daea;
    border-radius: 8px;
    padding: 16px 16px 12px;
    background: #f9faff;
    box-shadow: 0 2px 10px rgba(0,0,51,.07);
}
.chart-figure img {
    max-width: 100%; height: auto; display: block;
    margin: 0 auto 8px; border-radius: 4px;
    box-shadow: none;
}
.chart-figure figcaption {
    font-family: sans-serif; font-size: 0.88em; color: #444;
    margin-top: 6px; text-align: center; font-style: italic;
    line-height: 1.4;
}
/* standalone img (not in chart-figure) */
img         { max-width: 100%; height: auto; display: block; margin: 1.5em auto;
              border-radius: 4px; box-shadow: 0 2px 8px rgba(0,0,0,.12); }
/* ── Layout helpers ──────────────────────────────────────────────────── */
.toc        { background: #f5f9ff; border: 1px solid #b8d0f0; border-radius: 8px;
              padding: 1.2rem 1.8rem; margin: 2.5rem 0; font-family: sans-serif; }
.toc h2     { border: none; margin-top: 0; color: #003366; font-size: 1.1em; }
.toc ol     { margin: 0.4em 0; padding-left: 1.4em; }
.toc li     { margin: 0.25em 0; }
.toc a      { color: #003366; }
.title-page { text-align: center; padding: 70px 0 50px;
              border-bottom: 2px solid #dde; margin-bottom: 3.5rem; }
.title-page h1 { border: none; font-size: 2em; color: #002244; }
.section    { margin-bottom: 3.5rem; border-bottom: 1px solid #e0e4ef;
              padding-bottom: 1.5rem; }
/* Bibliography */
ol.bib-list { font-family: sans-serif; font-size: 0.92em; line-height: 1.6;
              padding-left: 1.6em; }
ol.bib-list li { margin-bottom: 0.6em; }
/* Source figures (scraped from PDFs/web, placed inline by RAG matching) */
.source-figure figure { background: #f5f7fa; border-color: #c8d6ea; }
.source-figure figcaption { font-size: 0.84em; color: #555; }
"""

# ── Lightbox CSS ───────────────────────────────────────────────────────────────
_LIGHTBOX_CSS = """
/* ── Zoom cursor on clickable elements ──────────────────────────────────── */
figure img, .section table {
  cursor: zoom-in;
  transition: opacity 0.15s, box-shadow 0.15s;
}
figure img:hover {
  opacity: 0.88;
  box-shadow: 0 4px 24px rgba(0,51,102,0.28);
}
.section table:hover {
  box-shadow: 0 4px 24px rgba(0,51,102,0.18);
}

/* ── Overlay backdrop ────────────────────────────────────────────────────── */
#lb-overlay {
  display: none;
  position: fixed;
  inset: 0;
  background: rgba(8, 12, 26, 0.90);
  backdrop-filter: blur(5px);
  -webkit-backdrop-filter: blur(5px);
  z-index: 9000;
  flex-direction: column;
  align-items: center;
  justify-content: center;
}
#lb-overlay.lb-open { display: flex; animation: lb-fadein 0.18s ease; }
@keyframes lb-fadein { from { opacity:0 } to { opacity:1 } }

/* ── Zoomed image ────────────────────────────────────────────────────────── */
#lb-img {
  max-width: 90vw;
  max-height: 78vh;
  object-fit: contain;
  border-radius: 5px;
  box-shadow: 0 10px 50px rgba(0,0,0,0.7);
  transform-origin: center center;
  transition: transform 0.1s ease;
  user-select: none;
  display: block;
}

/* ── Zoomed table wrapper ────────────────────────────────────────────────── */
#lb-table-wrap {
  max-width: 90vw;
  max-height: 78vh;
  overflow: auto;
  background: #fff;
  border-radius: 6px;
  padding: 1.4em 1.8em;
  box-shadow: 0 10px 50px rgba(0,0,0,0.7);
  display: none;
}
#lb-table-wrap table  { width: auto; min-width: 20em; border-collapse: collapse; }
#lb-table-wrap th     { background: #003366; color: #fff; padding: 6px 12px; }
#lb-table-wrap td     { border: 1px solid #ccc; padding: 5px 12px; font-size: 0.95em; }
#lb-table-wrap tr:nth-child(even) td { background: #f4f7fc; }

/* ── Caption + counter ───────────────────────────────────────────────────── */
#lb-caption {
  color: #dde3f4;
  font-family: sans-serif;
  font-size: 0.92em;
  margin-top: 0.8em;
  text-align: center;
  max-width: 80vw;
  line-height: 1.45;
  text-shadow: 0 1px 4px rgba(0,0,0,0.6);
}
#lb-counter {
  color: rgba(180,195,230,0.65);
  font-family: monospace;
  font-size: 0.8em;
  margin-top: 0.3em;
}

/* ── Toolbar (top bar) ───────────────────────────────────────────────────── */
#lb-toolbar {
  position: fixed;
  top: 0; left: 0; right: 0;
  display: flex;
  justify-content: space-between;
  align-items: center;
  padding: 0.55em 1.1em;
  z-index: 9010;
  pointer-events: none;
  background: linear-gradient(to bottom, rgba(0,0,0,0.45), transparent);
}
#lb-toolbar .lb-btn-group { display: flex; gap: 0.35em; pointer-events: all; }
#lb-toolbar button {
  pointer-events: all;
  background: rgba(255,255,255,0.13);
  border: 1px solid rgba(255,255,255,0.28);
  color: #fff;
  font-size: 1em;
  padding: 0.28em 0.7em;
  border-radius: 4px;
  cursor: pointer;
  transition: background 0.14s;
  line-height: 1.4;
}
#lb-toolbar button:hover { background: rgba(255,255,255,0.26); }

/* ── Prev / Next arrows ──────────────────────────────────────────────────── */
.lb-nav {
  position: fixed;
  top: 50%; transform: translateY(-50%);
  background: rgba(255,255,255,0.12);
  border: 1px solid rgba(255,255,255,0.28);
  color: #fff;
  font-size: 2em;
  line-height: 1;
  padding: 0.15em 0.45em;
  border-radius: 6px;
  cursor: pointer;
  z-index: 9010;
  user-select: none;
  transition: background 0.14s;
}
.lb-nav:hover { background: rgba(255,255,255,0.28); }
#lb-prev { left: 0.8em; }
#lb-next { right: 0.8em; }
.lb-nav.lb-hidden { opacity: 0; pointer-events: none; }

/* ── Zoom % badge (bottom-right, auto-hides) ─────────────────────────────── */
#lb-zoom-badge {
  position: fixed;
  bottom: 2.2em; right: 1em;
  background: rgba(255,255,255,0.14);
  color: #ccd;
  font-family: monospace;
  font-size: 0.8em;
  padding: 0.2em 0.6em;
  border-radius: 12px;
  z-index: 9010;
  pointer-events: none;
  opacity: 0;
  transition: opacity 0.3s;
}
#lb-zoom-badge.lb-show { opacity: 1; }

/* ── Keyboard hint strip (bottom-centre) ─────────────────────────────────── */
#lb-hints {
  position: fixed;
  bottom: 0.5em; left: 50%; transform: translateX(-50%);
  color: rgba(160,175,210,0.50);
  font-family: monospace;
  font-size: 0.72em;
  z-index: 9010;
  pointer-events: none;
  white-space: nowrap;
}
"""

# ── Lightbox JavaScript ────────────────────────────────────────────────────────
_LIGHTBOX_SCRIPT = """\
<script>
(function () {
  'use strict';

  /* ── State ─────────────────────────────────────────────────────────── */
  var items = [];          /* [{type:'img'|'table', el, caption}] */
  var current = -1;
  var scale = 1.0;
  var translateX = 0, translateY = 0;
  var dragStart = null;
  var zoomTimer = null;
  var MIN_SCALE = 0.4, MAX_SCALE = 8.0;

  /* ── Build overlay DOM ─────────────────────────────────────────────── */
  function buildUI() {
    var overlay = document.createElement('div');
    overlay.id = 'lb-overlay';
    overlay.innerHTML =
      '<div id="lb-toolbar">' +
        '<div class="lb-btn-group">' +
          '<button id="lb-zoom-out" title="Zoom out (-)">&#8722;</button>' +
          '<button id="lb-zoom-reset" title="Reset zoom (0)">1:1</button>' +
          '<button id="lb-zoom-in" title="Zoom in (+)">+</button>' +
        '</div>' +
        '<button id="lb-close" title="Close (Esc)">&#x2715;</button>' +
      '</div>' +
      '<button class="lb-nav" id="lb-prev">&#8249;</button>' +
      '<div id="lb-content">' +
        '<img id="lb-img" alt="">' +
        '<div id="lb-table-wrap"></div>' +
      '</div>' +
      '<div id="lb-caption"></div>' +
      '<div id="lb-counter"></div>' +
      '<button class="lb-nav" id="lb-next">&#8250;</button>' +
      '<div id="lb-zoom-badge">100%</div>' +
      '<div id="lb-hints">Esc&nbsp;close &nbsp;&#8592;&nbsp;&#8594;&nbsp;navigate &nbsp;scroll&nbsp;/&nbsp;+&#8722;&nbsp;zoom &nbsp;drag&nbsp;to&nbsp;pan</div>';
    document.body.appendChild(overlay);

    overlay.addEventListener('click', function (e) {
      if (e.target === overlay) closeLB();
    });
    g('lb-close').addEventListener('click', closeLB);
    g('lb-prev').addEventListener('click', function (e) { e.stopPropagation(); navigate(-1); });
    g('lb-next').addEventListener('click', function (e) { e.stopPropagation(); navigate(1); });
    g('lb-zoom-in').addEventListener('click', function (e) { e.stopPropagation(); applyZoom(0.25); });
    g('lb-zoom-out').addEventListener('click', function (e) { e.stopPropagation(); applyZoom(-0.25); });
    g('lb-zoom-reset').addEventListener('click', function (e) { e.stopPropagation(); resetView(); });

    /* Scroll-to-zoom on image */
    g('lb-img').addEventListener('wheel', function (e) {
      e.preventDefault();
      applyZoom(e.deltaY < 0 ? 0.18 : -0.18);
    }, { passive: false });

    /* Drag-to-pan */
    g('lb-img').addEventListener('mousedown', startDrag);
    window.addEventListener('mousemove', onDrag);
    window.addEventListener('mouseup', endDrag);

    /* Touch pinch-to-zoom */
    var lastPinchDist = 0;
    overlay.addEventListener('touchstart', function (e) {
      if (e.touches.length === 2) {
        lastPinchDist = pinchDist(e.touches);
      }
    }, { passive: true });
    overlay.addEventListener('touchmove', function (e) {
      if (e.touches.length === 2) {
        var d = pinchDist(e.touches);
        applyZoom((d - lastPinchDist) * 0.009);
        lastPinchDist = d;
      }
    }, { passive: true });
  }

  function g(id) { return document.getElementById(id); }
  function pinchDist(touches) {
    return Math.hypot(
      touches[0].clientX - touches[1].clientX,
      touches[0].clientY - touches[1].clientY
    );
  }

  /* ── Open / close ──────────────────────────────────────────────────── */
  function openLB(idx) {
    current = idx;
    resetView();
    render();
    g('lb-overlay').classList.add('lb-open');
    document.body.style.overflow = 'hidden';
  }

  function closeLB() {
    g('lb-overlay').classList.remove('lb-open');
    document.body.style.overflow = '';
    current = -1;
  }

  function navigate(dir) {
    if (!items.length) return;
    current = (current + dir + items.length) % items.length;
    resetView();
    render();
  }

  /* ── Render current item into modal ────────────────────────────────── */
  function render() {
    if (current < 0 || current >= items.length) return;
    var item = items[current];
    var imgEl = g('lb-img');
    var wrap  = g('lb-table-wrap');

    if (item.type === 'img') {
      imgEl.style.display = 'block';
      wrap.style.display = 'none';
      imgEl.src = item.el.src;
    } else {
      imgEl.style.display = 'none';
      wrap.style.display = 'block';
      wrap.innerHTML = item.el.outerHTML;
    }

    g('lb-caption').textContent = item.caption;
    g('lb-counter').textContent = (current + 1) + ' / ' + items.length;

    var hidden = items.length <= 1;
    g('lb-prev').classList.toggle('lb-hidden', hidden);
    g('lb-next').classList.toggle('lb-hidden', hidden);
  }

  /* ── Zoom & pan ────────────────────────────────────────────────────── */
  function applyZoom(delta) {
    scale = Math.min(MAX_SCALE, Math.max(MIN_SCALE, scale + delta));
    updateTransform();
    showZoomBadge();
  }

  function resetView() {
    scale = 1.0; translateX = 0; translateY = 0;
    updateTransform();
  }

  function updateTransform() {
    var img = g('lb-img');
    if (img) {
      img.style.transform =
        'translate(' + translateX + 'px, ' + translateY + 'px) scale(' + scale + ')';
      img.style.cursor = scale > 1.02 ? 'grab' : '';
    }
  }

  function showZoomBadge() {
    var badge = g('lb-zoom-badge');
    badge.textContent = Math.round(scale * 100) + '%';
    badge.classList.add('lb-show');
    clearTimeout(zoomTimer);
    zoomTimer = setTimeout(function () { badge.classList.remove('lb-show'); }, 1400);
  }

  function startDrag(e) {
    if (scale <= 1.02) return;
    dragStart = { x: e.clientX - translateX, y: e.clientY - translateY };
    g('lb-img').style.cursor = 'grabbing';
  }
  function onDrag(e) {
    if (!dragStart) return;
    translateX = e.clientX - dragStart.x;
    translateY = e.clientY - dragStart.y;
    updateTransform();
  }
  function endDrag() {
    dragStart = null;
    updateTransform();   /* restores grab/default cursor */
  }

  /* ── Keyboard ──────────────────────────────────────────────────────── */
  document.addEventListener('keydown', function (e) {
    if (!g('lb-overlay').classList.contains('lb-open')) return;
    switch (e.key) {
      case 'Escape':     closeLB();        break;
      case 'ArrowLeft':  navigate(-1);     break;
      case 'ArrowRight': navigate(1);      break;
      case '+': case '=': applyZoom(0.25); break;
      case '-':           applyZoom(-0.25);break;
      case '0':           resetView();     break;
    }
  });

  /* ── Collect items & attach click handlers ─────────────────────────── */
  function init() {
    buildUI();

    /* Images inside <figure> elements */
    document.querySelectorAll('figure img').forEach(function (img) {
      var fig = img.closest('figure');
      var capEl = fig && fig.querySelector('figcaption');
      var cap = capEl ? capEl.textContent.trim() : (img.alt || '');
      var idx = items.length;
      items.push({ type: 'img', el: img, caption: cap });
      img.addEventListener('click', function () { openLB(idx); });
    });

    /* Data tables inside sections */
    document.querySelectorAll('.section table').forEach(function (tbl) {
      var cap = '';
      var sib = tbl.previousElementSibling;
      while (sib) {
        if (/^H[1-6]$/.test(sib.tagName)) { cap = sib.textContent.trim(); break; }
        if (sib.tagName === 'P') { cap = sib.textContent.trim().slice(0, 150); break; }
        sib = sib.previousElementSibling;
      }
      if (!cap) cap = 'Table';
      var idx = items.length;
      items.push({ type: 'table', el: tbl, caption: cap });
      tbl.addEventListener('click', function (e) {
        if (e.target.tagName === 'A') return;   /* don't intercept link clicks */
        openLB(idx);
      });
    });
  }

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', init);
  } else {
    init();
  }
}());
</script>"""


@tool
def build_html(
    sections: list[dict],
    figures: list[dict],
    bibliography: list[dict],
    output_path: str = "",
    include_toc: bool = True,
    title: str = "ASTRA Research Report",
) -> dict:
    """
    Layer 6: Build a self-contained HTML report with:
    - MathJax for LaTeX formula rendering ($...$ and $$...$$)
    - Active hyperlinks from markdown [text](url) syntax
    - Responsive design with syntax-highlighted code blocks
    - Embedded chart images

    Args:
        sections: List of {title, markdown, citations, chart_paths} dicts.
        figures: List of {path, caption} dicts.
        bibliography: List of {key, title, authors, year, url} dicts.
        output_path: Output .html file path.
        include_toc: Include table of contents.
        title: Report title.

    Returns:
        {html_path, word_count}
    """
    import time
    import base64

    cfg = get_config()
    if not output_path:
        ts = int(time.time())
        output_path = str(cfg.get_output_path(f"ASTRA_Report_{ts}.html"))

    logger.info(f"[Layer 6] Building HTML: {output_path}")

    try:
        parts: list[str] = []
        total_words = 0
        fig_counter = [0]  # mutable so nested helper can increment

        # Title page
        parts.append(
            '<div class="title-page">'
            f"<h1>{title}</h1>"
            '<p style="color:#666;font-size:1.05em;font-family:sans-serif">'
            "Generated by ASTRA · Multi-Agent Research Pipeline</p>"
            "</div>"
        )

        # Table of contents
        if include_toc:
            toc_items = "".join(
                f'<li><a href="#sec-{i}">{sec.get("title", f"Section {i}")}</a></li>'
                for i, sec in enumerate(sections, 1)
            )
            parts.append(
                '<div class="toc">'
                "<h2>Table of Contents</h2>"
                f"<ol>{toc_items}</ol>"
                "</div>"
            )

        # Sections
        for i, sec in enumerate(sections, 1):
            title_str = sec.get("title", f"Section {i}")
            content = sec.get("markdown", "")
            chart_paths = sec.get("chart_paths", [])

            section_html = _markdown_to_html(content)

            # Link [Source N] citations to bibliography anchors
            section_html = re.sub(
                r"\[Source (\d+)\]",
                r'<a href="#bib-\1" class="cite">[Source \1]</a>',
                section_html,
            )

            # Remove redundant per-section References/Key Sources subsections
            section_html = _strip_section_references(section_html)

            # Make bare URLs in text nodes clickable
            section_html = _linkify_bare_urls(section_html)

            total_words += len(content.split())

            # ── Inline chart injection (ASTRA-generated from data tables) ──
            # Each chart is a visualisation of a table in this section.
            # Inject directly after the last </table> — scientific-paper style.
            for cp in chart_paths:
                if not cp or not Path(cp).exists():
                    continue
                fig_counter[0] += 1
                src = _embed_image(cp)
                chart_fig_html = (
                    f'\n<div class="chart-figure">'
                    f"<figure>"
                    f'<img src="{src}" alt="Figure {fig_counter[0]}: {title_str}">'
                    f"<figcaption><strong>Figure {fig_counter[0]}:</strong> "
                    f"{title_str}</figcaption>"
                    f"</figure></div>\n"
                )
                table_end_pos = section_html.rfind("</table>")
                if table_end_pos >= 0:
                    insert_at = table_end_pos + len("</table>")
                    section_html = (
                        section_html[:insert_at]
                        + chart_fig_html
                        + section_html[insert_at:]
                    )
                else:
                    section_html += chart_fig_html

            # ── Inline source-figure injection (scraped, RAG-matched) ─────
            # Figures were selected per-section by figure_search() in
            # node_draft_report using bge-m3 cosine similarity, so every
            # figure here is topically relevant to this section.
            # We woven them into the narrative immediately after the paragraph
            # whose text best overlaps the figure's description keywords.
            for sfig in sec.get("source_figures", []):
                fp = sfig.get("path", "")
                if not fp or not Path(fp).exists():
                    continue
                # Guard: skip figures that were de-ranked below threshold
                # (shouldn't happen after figure_search, but defensive)
                if sfig.get("score", 1.0) < 0.20:
                    continue
                cap = sfig.get("caption", "") or sfig.get("title", "")
                # Skip tiny images (icons, blank pages) — file size proxy
                try:
                    if Path(fp).stat().st_size < 8_000:   # < 8 KB
                        continue
                except OSError:
                    continue
                fig_counter[0] += 1
                src = _embed_image(fp)
                display_cap = cap or f"Source Figure {fig_counter[0]}"
                sfig_html = (
                    f'\n<div class="chart-figure source-figure">'
                    f"<figure>"
                    f'<img src="{src}" alt="Figure {fig_counter[0]}: '
                    f'{display_cap[:80]}">'
                    f"<figcaption><strong>Figure {fig_counter[0]}:</strong> "
                    f"{display_cap}</figcaption>"
                    f"</figure></div>\n"
                )
                # Use the full RAG chunk text (VLM description when available,
                # otherwise just the caption) for keyword matching so we find
                # the most relevant paragraph even without VLM output.
                match_text = sfig.get("description") or cap
                insert_at = _best_injection_point(section_html, match_text)
                section_html = (
                    section_html[:insert_at] + sfig_html + section_html[insert_at:]
                )

            parts.append(
                f'<div class="section" id="sec-{i}">'
                f"<h1>{title_str}</h1>"
                f"{section_html}"
                "</div>"
            )

        # Additional standalone figures (not tied to a specific section)
        if figures:
            parts.append('<div class="section"><h1>Figures</h1>')
            for fig in figures:
                fp = fig.get("path", "")
                if fp and Path(fp).exists():
                    fig_counter[0] += 1
                    src = _embed_image(fp)
                    cap = fig.get("caption", "")
                    parts.append(
                        f'<div class="chart-figure"><figure>'
                        f'<img src="{src}" alt="{cap}">'
                        f"<figcaption><strong>Figure {fig_counter[0]}:</strong> "
                        f"{cap}</figcaption></figure></div>"
                    )
            parts.append("</div>")

        # Bibliography — id="bib-N" enables [Source N] anchor links from sections
        bib_items = ""
        for j, bib in enumerate(bibliography, 1):
            authors = ", ".join(bib.get("authors", [bib.get("key", "Unknown")]))
            year = bib.get("year", "n.d.")
            title_b = bib.get("title", "")
            url = bib.get("url", "")
            url_link = (
                f'<a href="{url}" target="_blank">{url}</a>' if url else ""
            )
            bib_items += (
                f'<li id="bib-{j}">{authors} ({year}). '
                f"<em>{title_b}</em>. {url_link}</li>"
            )
        parts.append(f'<h1>Bibliography</h1><ol class="bib-list">{bib_items}</ol>')

        full_html = (
            "<!DOCTYPE html>\n<html lang='en'>\n<head>\n"
            '<meta charset="utf-8">\n'
            '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
            f"<title>{title}</title>\n"
            f"<style>{_HTML_EXTRA_CSS}\n{_LIGHTBOX_CSS}</style>\n"
            f"{_MATHJAX_SCRIPT}\n"
            "</head>\n<body>\n"
            + "\n".join(parts)
            + f"\n{_LIGHTBOX_SCRIPT}\n</body>\n</html>"
        )

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_text(full_html, encoding="utf-8")

        logger.info(f"[Layer 6] HTML built: {output_path} ({total_words}w, self-contained)")
        return {"html_path": output_path, "word_count": total_words}

    except Exception as e:
        logger.error(f"[Layer 6] build_html failed: {e}")
        return {"html_path": "", "word_count": 0, "error": str(e)}


# ─── embed_figure ────────────────────────────────────────────────────────────

@tool
def embed_figure(
    markdown_content: str,
    figure_path: str,
    caption: str = "",
    after_keyword: str = "",
) -> str:
    """
    Layer 6: Embed a figure reference into markdown content.

    Args:
        markdown_content: The markdown text to modify.
        figure_path: Path to the figure image.
        caption: Figure caption text.
        after_keyword: Insert figure after the first occurrence of this keyword.

    Returns:
        Updated markdown string with figure embedded.
    """
    figure_md = f"\n\n![{caption}]({figure_path})\n*Figure: {caption}*\n\n"

    if after_keyword and after_keyword in markdown_content:
        idx = markdown_content.index(after_keyword) + len(after_keyword)
        # Insert after the current paragraph
        next_para = markdown_content.find("\n\n", idx)
        if next_para > 0:
            return markdown_content[:next_para] + figure_md + markdown_content[next_para:]

    return markdown_content + figure_md
